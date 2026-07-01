#!/usr/bin/env python3
"""
app.py — Redrob Hackathon: Intelligent Candidate Discovery & Ranking (v3)

Self-contained Gradio app for HuggingFace Spaces.
All scoring/ranking logic from rank.py is embedded directly below — no external
rank.py import, no hardcoded candidate IDs/ranks/scores/company names.

Upload a candidates JSONL(.gz) file + a job description (.md/.txt/.docx, or
paste JD text), click Run, and download the ranked submission.csv.
"""

import csv
import gzip
import hashlib
import io
import json
import os
import re
import sys
import tempfile
import time
from collections import namedtuple
from datetime import date, datetime
from pathlib import Path

os.environ["TRANSFORMERS_OFFLINE"] = "0"
os.environ["HF_DATASETS_OFFLINE"] = "0"
os.environ["HF_HUB_OFFLINE"] = "0"

import faiss
import gradio as gr
import numpy as np
import pandas as pd
from sentence_transformers import SentenceTransformer

TODAY = date.today()
BIENCODER_MODEL = "BAAI/bge-small-en-v1.5"

TIEBREAK_DELTA = 0.002

W_SEMANTIC = 0.30
W_RULES = 0.70

W_SKILL = 0.30
W_CAREER = 0.35
W_LOC = 0.10
W_BEHAV = 0.25

assert abs(W_SKILL + W_CAREER + W_LOC + W_BEHAV - 1.0) < 1e-9

RuleScores = namedtuple("RuleScores", ["total", "skill", "career", "loc", "behav"])


# ══════════════════════════════════════════════════════════════════════════════
# SKILL REFERENCE SETS
# ══════════════════════════════════════════════════════════════════════════════

SKILLS_CORE_JD = frozenset({
    "retrieval", "retrieval pipeline", "retrieval pipelines", "dense retrieval",
    "ranking", "ranking systems", "learning to rank", "ltr", "learning-to-rank",
    "reranking", "embeddings", "text embeddings", "embedding pipeline",
    "vector search", "vector database", "vector index", "semantic search",
    "faiss", "pinecone", "qdrant", "weaviate", "milvus", "opensearch",
    "elasticsearch", "chroma",
    "production ml", "mlops", "python", "bm25",
    "hybrid search", "neural search",
    "ndcg", "mrr", "map", "recall@k", "precision@k",
    "search systems", "search pipeline", "search engine",
    "recommendation system", "recommendation engine",
})

SKILLS_SUPPORTING = frozenset({
    "langchain", "llamaindex", "llama index", "prompt engineering",
    "lora", "qlora", "peft", "fine-tuning", "fine-tuning llms",
    "openai api", "chatgpt", "gpt-4", "gpt4",
    "langsmith", "langgraph",
})

SKILLS_TIER_A = frozenset({
    "embeddings", "sentence-transformers", "sentence transformers",
    "faiss", "pinecone", "qdrant", "weaviate", "milvus", "chroma",
    "opensearch", "elasticsearch", "vector search", "hybrid search",
    "semantic search", "dense retrieval", "information retrieval",
    "neural search", "vector index", "embedding pipeline",
    "ndcg", "mrr", "map", "reranking", "learning to rank", "ltr",
    "learning-to-rank", "ranking systems", "ai ranking", "intelligent ranking",
    "recall@k", "precision@k", "ranking evaluation", "retrieval evaluation",
    "a/b testing", "a/b test", "ab testing", "offline evaluation", "online evaluation",
    "offline-to-online", "evaluation framework",
    "retrieval quality", "retrieval quality regression", "ranking regression",
    "embedding drift", "index refresh", "index rebuild",
    "retrieval-quality regression",
    "nlp", "llm", "large language model", "transformers", "bert",
    "rag", "retrieval augmented generation", "text embeddings",
    "openai embeddings", "bge", "e5",
    "llm workflows", "llm powered", "llm-powered",
    "python", "mlops", "production ml", "ranking", "bm25",
    "retrieval pipeline", "retrieval pipelines",
    "search pipeline", "search systems",
    "recommendation system", "recommendation engine",
    "candidate ranking", "talent intelligence", "candidate discovery",
    "intelligent recommendation",
    "search engine", "relevance engineering", "search relevance",
    "query understanding", "query expansion", "document ranking",
    "two-tower model", "dual encoder", "cross encoder", "bi-encoder",
    "approximate nearest neighbor", "ann", "knn search",
    "vector database", "inverted index",
    "hybrid retrieval", "sparse retrieval", "colbert", "splade",
    "sentence transformers", "cross-encoder", "bi-encoder",
    "query rewriting", "query classification",
    "production monitoring", "model monitoring", "data drift",
})

SKILLS_TIER_B = frozenset({
    "lora", "qlora", "peft", "fine-tuning", "fine-tuning llms",
    "xgboost", "lightgbm", "tfidf", "neural ranking",
    "learning to rank models",
    "kafka", "spark", "fastapi", "bentoml", "triton",
    "model serving", "inference optimization", "distributed systems",
    "large-scale inference", "large scale inference",
    "hugging face", "huggingface", "pytorch", "tensorflow", "scikit-learn",
    "aws", "gcp", "azure", "docker", "kubernetes",
    "feature engineering", "data pipelines",
    "experimentation", "online evaluation",
    "hr-tech", "hr tech", "recruiting tech", "recruitech",
    "marketplace", "marketplace products", "talent marketplace",
    "open-source", "open source", "open source contribution",
    "open source contributions",
    "grpc", "rest api", "redis", "celery", "airflow", "mlflow",
    "feature store", "online store", "model registry",
    "langchain", "llamaindex", "llama index", "prompt engineering",
    "langsmith", "langgraph",
})

SKILLS_PENALTY = frozenset({
    "computer vision", "image classification", "object detection",
    "speech recognition", "asr", "tts", "text to speech",
    "robotics", "solidworks", "ansys", "autocad",
    "sap", "six sigma", "photoshop", "illustrator", "figma",
    "marketing", "seo", "content writing", "accounting",
    "crm", "salesforce", "tableau", "power bi",
})

PRODUCTION_KW = frozenset({
    "production", "deployed", "shipped", "serving", "at scale",
    "real users", "latency", "inference", "live system",
    "millions", "query per second", "qps", "p99", "sla", "uptime",
    "embedding drift", "index refresh", "retrieval quality",
    "production ml", "production system", "production environment",
    "online inference", "serving models", "model serving",
    "inference api", "inference pipeline", "real-time inference",
    "low latency", "high throughput", "traffic", "requests per second",
    "billion", "hundred million", "scale to",
    "kubernetes", "docker", "k8s", "helm", "containerized",
    "fastapi", "flask api", "grpc service",
    "a/b test", "shadow mode", "canary deploy", "blue-green",
    "monitoring", "alerting", "observability",
    "feature store", "online feature", "real-time feature",
    "api endpoint", "rest endpoint", "microservice", "service mesh",
    "load balancer", "autoscaling", "ci/cd", "deployment pipeline",
    "production deployment", "model deployment", "model serving",
    "online serving", "batch inference", "streaming inference",
    "production traffic", "live traffic", "production environment",
    "rollout", "gradual rollout", "model rollout",
    "production ml system", "ml system", "ai system",
    "search service", "retrieval service", "ranking service",
    "recommendation service", "embedding service",
    "latency optimization", "latency optimisation",
    "throughput optimization", "memory optimization",
    "production monitoring", "model monitoring", "drift detection",
    "index rebuild", "index update", "index refresh",
})

PROD_RETRIEVAL_KW = frozenset({
    "search engine", "retrieval system", "retrieval pipeline",
    "recommendation engine", "recommendation system",
    "ranking system", "ranking pipeline",
    "semantic search", "vector search", "faiss", "pinecone", "qdrant",
    "weaviate", "milvus", "elasticsearch", "opensearch", "solr",
    "embedding pipeline", "vector database", "vector index",
    "neural search", "dense retrieval", "sparse retrieval",
    "search infrastructure", "search platform", "search ranking",
    "query serving", "index serving", "ann index",
    "two-tower", "dual encoder", "cross encoder",
    "reranker", "reranking", "learning to rank",
    "ndcg", "mrr", "map", "recall@k", "precision@k",
    "hybrid retrieval", "bm25", "colbert", "splade",
    "search relevance", "relevance engineering",
    "query understanding", "query rewriting",
    "document ranking", "passage ranking",
    "embedding search", "knn search", "ann search",
    "production search", "production retrieval",
    "production recommendation", "production ranking",
})

PRODUCT_ENG_KW = frozenset({
    "owned", "ownership", "end-to-end", "built from scratch",
    "led", "designed and built", "architected", "scaled",
    "search infrastructure", "recommendation infrastructure",
    "ranking infrastructure", "retrieval infrastructure",
    "platform engineering", "ml platform", "ai platform",
    "search platform", "recommendation platform",
    "production ai", "production ml system", "production ai system",
    "zero to one", "0 to 1", "greenfield",
    "tech lead", "technical lead", "engineering lead",
    "led the team", "led development", "led engineering",
    "principal engineer", "staff engineer",
})

SENIORITY_SIGNALS = {
    "senior": 3, "staff": 5, "principal": 6, "lead": 4, "tech lead": 4,
    "technical lead": 4, "engineering lead": 4, "manager": 3, "director": 6,
    "vp": 7, "head of": 5, "architect": 5, "junior": 1, "associate": 2,
    "intern": 0, "fresher": 1,
}

ACADEMIC_ONLY_KW = frozenset({
    "research paper", "arxiv", "published paper", "academic project",
    "thesis", "dissertation", "coursework", "class project",
    "kaggle", "hackathon project", "side project", "personal project",
    "langchain demo", "openai api demo",
    "proof of concept", "poc",
})

PRODUCT_CO_SIGNALS = frozenset({
    "product", "platform", "saas", "b2b", "b2c", "marketplace",
    "our platform", "our product", "our app", "our service",
    "users", "customers", "dau", "mau", "retention", "conversion",
    "revenue", "growth", "monetization", "subscription",
    "shipped", "launched", "released", "v1", "v2", "milestone",
    "roadmap", "sprint", "agile", "scrum",
    "founding team", "startup", "series", "seed", "funded",
})

JD_EXACT_KW = frozenset({
    "embedding drift", "index refresh", "retrieval quality regression",
    "retrieval-quality regression", "a/b test", "offline-to-online",
    "ndcg", "mrr", "eval framework", "ranking regression",
    "retrieval quality", "index rebuild", "vector index",
})

OPS_SIGNALS: list[tuple[str, str]] = [
    ("embedding drift", "embedding drift monitoring"),
    ("index refresh", "index refresh"),
    ("ndcg", "NDCG evaluation"),
    ("mrr", "MRR evaluation"),
    ("a/b test", "A/B testing"),
    ("offline-to-online", "offline-to-online eval"),
    ("retrieval quality", "retrieval quality monitoring"),
    ("search engine", "search engine"),
    ("recommendation", "recommendation system"),
    ("vector search", "vector search"),
    ("semantic search", "semantic search"),
    ("reranking", "reranking pipeline"),
    ("learning to rank", "LTR"),
    ("latency", "latency optimisation"),
    ("production monitoring", "production monitoring"),
    ("bm25", "BM25"),
    ("hybrid retrieval", "hybrid retrieval"),
    ("colbert", "ColBERT"),
]

NOTABLE_TECH = [
    "faiss", "pinecone", "qdrant", "weaviate", "milvus", "chroma",
    "elasticsearch", "opensearch", "solr",
    "bm25", "colbert", "splade",
    "sentence transformers", "bge", "e5", "openai embeddings",
    "hybrid retrieval", "dense retrieval", "sparse retrieval",
    "two-tower", "dual encoder", "cross encoder",
    "ndcg", "mrr", "learning to rank", "ltr",
    "rag", "retrieval augmented generation",
    "fastapi", "kubernetes", "docker", "triton", "bentoml",
]

ACTION_VERBS = [
    "Built", "Designed", "Developed", "Implemented", "Delivered",
    "Productionized", "Scaled", "Maintained", "Led development of",
    "Shipped", "Deployed", "Architected", "Launched", "Engineered",
]

KNOWN_TECH_CITIES = [
    "pune", "noida", "hyderabad", "bangalore", "bengaluru", "mumbai",
    "delhi", "gurgaon", "gurugram", "ncr", "delhi ncr", "new delhi",
    "chennai", "kolkata", "ahmedabad", "kochi", "jaipur", "indore", "coimbatore",
    "new york", "san francisco", "london", "singapore", "dubai",
    "berlin", "amsterdam", "toronto", "sydney", "seattle", "boston",
]

PROFICIENCY_W: dict[str, float] = {
    "beginner": 0.25, "intermediate": 0.60, "advanced": 1.00, "expert": 1.20,
}

SERVICES_COMPANIES = frozenset({
    "tcs", "infosys", "wipro", "accenture", "cognizant", "capgemini",
    "hcl", "tech mahindra", "hexaware", "mphasis", "mindtree",
    "l&t infotech", "ltimindtree",
})

KNOWN_PRODUCT_COMPANIES = frozenset({
    "google", "meta", "microsoft", "amazon", "apple", "netflix", "uber",
    "airbnb", "linkedin", "twitter", "x.com", "openai", "anthropic",
    "deepmind", "nvidia", "salesforce", "adobe", "stripe", "spotify",
    "flipkart", "zomato", "swiggy", "paytm", "razorpay", "phonepe",
    "meesho", "cred", "urban company", "urbanclap", "dunzo", "zepto",
    "blinkit", "groww", "zerodha", "sharechat", "moj", "dailyhunt",
    "inmobi", "freshworks", "zoho", "browserstack", "chargebee",
    "postman", "hasura", "setu", "sarvam", "sarvam ai", "krutrim",
    "yellow.ai", "yellow ai", "uniphore", "observe.ai", "observe ai",
    "mad street den", "mad street", "niramai", "sigtuple", "healthifyme",
    "practo", "1mg", "pharmeasy", "cult.fit", "curefit", "lenskart",
    "nykaa", "mamaearth", "boat", "vedantu", "byju", "unacademy",
    "upgrad", "physics wallah", "scaler", "car dekho", "cardekho",
    "cars24", "spinny", "droom", "ola", "rapido", "porter",
    "licious", "country delight", "bigbasket", "jiomart",
    "juspay", "cashfree", "instamojo", "billdesk",
    "moengage", "clevertap", "webengage", "netcore",
    "darwinbox", "keka", "greythr", "sumhr",
    "leadsquared", "chargebee", "exotel", "knowlarity",
    "niki.ai", "haptik", "gupshup",
})

NON_TECH_HONEYPOT_TITLES = frozenset({
    "accountant", "hr manager", "marketing manager", "civil engineer",
    "mechanical engineer", "sales manager", "operations manager",
    "customer support", "content writer", "brand manager",
})

TECH_EVIDENCE_KW = frozenset({
    "python", "machine learning", "nlp", "vector", "embedding",
    "retrieval", "transformer", "deep learning", "model",
    "ranking", "recommendation", "search",
})


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PATTERNS
# ══════════════════════════════════════════════════════════════════════════════

ML_ADJACENT_TITLE_RE = re.compile(
    r'data scientist|senior data scientist|data science|'
    r'software engineer.*\bml\b|swe.*\bml\b|software engineer.*\bai\b|'
    r'senior software engineer',
    re.IGNORECASE,
)

TITLE_FIT_RE = re.compile(
    r'\bml\b|machine learning|ai engineer|\bnlp\b|'
    r'research engineer|applied (scientist|ml|ai)|search engineer|'
    r'retrieval|ranking engineer|recommendation|\bllm\b|generative ai|'
    r'deep learning|senior.*engineer|staff engineer|principal engineer|'
    r'embedding|intelligence engineer|ml engineer|applied scientist|'
    r'relevance engineer|platform.*engineer|search.*engineer|'
    r'recommendation.*engineer|backend.*search|nlp engineer|'
    r'software engineer.*retrieval|software engineer.*search|'
    r'software engineer.*recommendation|software engineer.*ranking',
    re.IGNORECASE,
)

TITLE_PENALTY_RE = re.compile(
    r'\bmarketing\b|operations manager|hr manager|accountant|'
    r'civil engineer|mechanical engineer|customer support|'
    r'\bsales\b|content writer|business analyst|financial analyst|'
    r'brand manager',
    re.IGNORECASE,
)

WRONG_DOMAIN_RE = re.compile(
    r'computer vision|cv engineer|\bimage\b|speech|robotics|mechanical|civil|hardware',
    re.IGNORECASE,
)

JUNIOR_TITLE_RE = re.compile(
    r'\bjunior\b|\bentry.level\b|\bfresher\b|\bintern\b|\bgraduate trainee\b',
    re.IGNORECASE,
)

RETRIEVAL_TITLE_RE = re.compile(
    r'search|retrieval|relevance|recommendation|ranking|platform.*ml|'
    r'platform.*ai|backend.*search|nlp engineer',
    re.IGNORECASE,
)

SENIOR_TITLE_RE = re.compile(
    r'\bstaff\b|\bprincipal\b|\btech lead\b|\btechnical lead\b|'
    r'\bengineering lead\b|\blead engineer\b|\bhead of\b|\bdirector\b|\bvp\b',
    re.IGNORECASE,
)

NON_ML_PRIMARY_TITLES = frozenset({
    "data scientist", "senior data scientist", "data science lead",
    "senior software engineer", "software engineer", "senior data engineer",
    "senior software engineer (ml)", "software engineer (ml)",
    "software engineer (ai)", "senior software engineer (ai)",
    "data engineer",
})

TECH_TITLE_KW = frozenset({
    "engineer", "scientist", "ml", "ai", "nlp", "research",
    "machine learning", "data sci", "deep learning", "llm",
    "python", "ranking", "retrieval", "search", "embedding",
    "relevance", "recommendation", "platform", "backend",
})


# ══════════════════════════════════════════════════════════════════════════════
# UTILITIES
# ══════════════════════════════════════════════════════════════════════════════

def parse_date(s: str | None) -> date | None:
    if not s:
        return None
    try:
        return datetime.strptime(s[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return None


def days_since(d: date | None) -> int:
    if d is None:
        return 9999
    return max(0, (TODAY - d).days)


def _safe_truncate(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    cut = text[:max_chars].rsplit(" ", 1)[0]
    return cut.rstrip(".,;") + "..."


def _count_hits(text: str, kw_set: frozenset) -> int:
    return sum(1 for kw in kw_set if kw in text)


def _count_production_hits(desc_lower: str) -> int:
    return _count_hits(desc_lower, PRODUCTION_KW)


def _count_prod_retrieval_hits(text_lower: str) -> int:
    return _count_hits(text_lower, PROD_RETRIEVAL_KW)


def _infer_production(desc_lower: str) -> bool:
    prod_hits = _count_production_hits(desc_lower)
    retrieval_hits = _count_prod_retrieval_hits(desc_lower)
    infra_hits = sum(1 for kw in ("kubernetes", "k8s", "docker", "fastapi", "grpc service", "microservice") if kw in desc_lower)
    scale_hits = sum(1 for kw in ("millions", "billion", "qps", "latency", "traffic", "at scale") if kw in desc_lower)
    return (
        prod_hits >= 2 or
        (prod_hits >= 1 and retrieval_hits >= 1) or
        (infra_hits >= 1 and scale_hits >= 1)
    )


def _is_product_company_job(job: dict) -> bool:
    co = job.get("company", "").lower()
    if any(s in co for s in SERVICES_COMPANIES):
        return False
    if any(p in co for p in KNOWN_PRODUCT_COMPANIES):
        return True
    text = (
        job.get("description", "") + " " +
        job.get("company", "") + " " +
        job.get("industry", "")
    ).lower()
    return _count_hits(text, PRODUCT_CO_SIGNALS) >= 2


def _academic_penalty(full_text: str) -> float:
    hits = _count_hits(full_text, ACADEMIC_ONLY_KW)
    if hits == 0:
        return 0.0
    if hits == 1:
        return -0.03
    if hits == 2:
        return -0.06
    return -0.10


def _product_engineering_bonus(full_text: str) -> float:
    hits = _count_hits(full_text, PRODUCT_ENG_KW)
    if hits >= 4:
        return 0.08
    if hits >= 2:
        return 0.05
    if hits >= 1:
        return 0.02
    return 0.0


def _career_progression_score(career: list[dict]) -> float:
    if not career or len(career) < 2:
        return 0.0

    def _title_level(title: str) -> int:
        tl = title.lower()
        level = 2
        for kw, lvl in SENIORITY_SIGNALS.items():
            if kw in tl:
                level = max(level, lvl)
        return level

    def _job_sort_key(j: dict) -> str:
        return j.get("start_date", "") or ""

    sorted_jobs = sorted(career, key=_job_sort_key, reverse=True)
    levels = [_title_level(j.get("title", "")) for j in sorted_jobs]

    if len(levels) < 2:
        return 0.0

    delta = levels[0] - levels[-1]
    if delta >= 2:
        return 0.05
    if delta == 1:
        return 0.03
    if delta == 0:
        return 0.0
    return -0.03


def _get_action_verb(candidate_id: str) -> str:
    idx = int(hashlib.md5(str(candidate_id).encode()).hexdigest(), 16) % len(ACTION_VERBS)
    return ACTION_VERBS[idx]


def _extract_notable_tech(full_text: str) -> list[str]:
    found = [t for t in NOTABLE_TECH if t in full_text]
    seen: set[str] = set()
    out: list[str] = []
    for t in found:
        key = t.split()[0]
        if key not in seen:
            seen.add(key)
            out.append(t)
        if len(out) >= 3:
            break
    return out


# ══════════════════════════════════════════════════════════════════════════════
# JD PARSER
# ══════════════════════════════════════════════════════════════════════════════

def parse_jd(jd_text: str) -> dict:
    t = jd_text.lower()

    m = re.search(
        r'(\d+)\s*(?:to|-|–)\s*(\d+)\s*years?|(\d+)\+?\s*years?\s*(?:of\s*)?experience',
        t,
    )
    if m:
        if m.group(1) and m.group(2):
            yoe_min, yoe_max = int(m.group(1)), int(m.group(2))
        else:
            yoe_min = int(m.group(3))
            yoe_max = yoe_min + 4
    else:
        if any(w in t for w in ("senior", "staff", "lead", "principal")):
            yoe_min, yoe_max = 5, 12
        elif any(w in t for w in ("junior", "entry", "fresher", "graduate")):
            yoe_min, yoe_max = 0, 3
        else:
            yoe_min, yoe_max = 3, 8

    target_cities = [city for city in KNOWN_TECH_CITIES if city in t]

    country_hints: dict[str, list[str]] = {
        "india": ["india", "indian", "inr", "lpa", "lakhs",
                  "bengaluru", "hyderabad", "pune", "noida"],
        "usa": ["united states", " usa", "usd", "silicon valley"],
        "uk": ["united kingdom", " uk ", "gbp", "london"],
        "singapore": ["singapore", "sgd"],
        "uae": ["dubai", "aed", "uae"],
    }
    target_country: str | None = None
    for country, hints in country_hints.items():
        if any(h in t for h in hints):
            target_country = country
            break

    jd_skills = {s for s in (SKILLS_TIER_A | SKILLS_TIER_B) if s in t}

    seniority = "mid"
    for level, kws in (
        ("staff", ("staff engineer", "principal engineer", "architect role")),
        ("senior", ("senior", "sr.", "lead", "founding", "principal")),
        ("junior", ("junior", "entry level", "fresher")),
    ):
        if any(k in t for k in kws):
            seniority = level
            break

    if "fully remote" in t or "100% remote" in t:
        work_mode = "remote"
    elif "onsite" in t or "on-site" in t:
        work_mode = "onsite"
    elif "hybrid" in t:
        work_mode = "hybrid"
    elif "remote" in t:
        work_mode = "remote"
    else:
        work_mode = "hybrid"

    prefers_product = any(w in t for w in (
        "product company", "founding team", "startup",
        "not consulting", "not it services", "product-first",
        "series a", "series b", "consulting firms", "we're building",
        "we are building", "not a fit", "bad fit",
    ))

    effective_tier_a = SKILLS_TIER_A | jd_skills

    return {
        "raw_text": jd_text,
        "yoe_min": yoe_min,
        "yoe_max": yoe_max,
        "target_cities": target_cities,
        "target_country": target_country,
        "jd_skills": jd_skills,
        "seniority": seniority,
        "work_mode": work_mode,
        "prefers_product": prefers_product,
        "_effective_tier_a": effective_tier_a,
    }


def _is_binary_garbage(text: str) -> bool:
    raw = text[:16]
    if raw.startswith("PK"):
        return True
    if raw.startswith("%PDF"):
        return True
    sample = text[:256]
    non_print = sum(1 for c in sample if ord(c) < 32 and c not in "\n\r\t")
    return non_print / max(len(sample), 1) > 0.10


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception:
        import zipfile, re as _re
        with zipfile.ZipFile(str(path)) as zf:
            with zf.open("word/document.xml") as xf:
                xml = xf.read().decode("utf-8", errors="replace")
        text = _re.sub(r"<[^>]+>", " ", xml)
        text = _re.sub(r"\s+", " ", text).strip()
        return text


def load_jd_from_path(jd_path: str, log: list[str]) -> str:
    p = Path(jd_path)
    if not p.exists():
        raise gr.Error(f"JD file not found: {jd_path}")

    magic = p.read_bytes()[:4]
    is_zip_container = magic[:2] == b"PK"
    is_pdf = magic[:4] == b"%PDF"

    text: str | None = None
    used_encoding = "unknown"

    if is_zip_container:
        log.append(f"Detected .docx/ZIP format for {p.name} — extracting text...")
        text = _extract_docx(p)
        used_encoding = "docx"
    elif is_pdf:
        raise gr.Error(f"{p.name} is a PDF. Convert to .docx or .md first.")
    else:
        for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
            try:
                candidate_text = p.read_text(encoding=enc)
                if _is_binary_garbage(candidate_text):
                    log.append(f"[{enc}] decoded to binary garbage — treating as .docx ZIP...")
                    text = _extract_docx(p)
                    used_encoding = "docx-fallback"
                    break
                text = candidate_text
                used_encoding = enc
                break
            except (UnicodeDecodeError, LookupError):
                continue

    if not text or not text.strip():
        raise gr.Error(f"Could not extract text from {jd_path}.")

    log.append(f"Loaded JD: {p.name} ({len(text.split())} words) [format: {used_encoding}]")
    return text


# ══════════════════════════════════════════════════════════════════════════════
# HONEYPOT DETECTION
# ══════════════════════════════════════════════════════════════════════════════

def is_honeypot(c: dict) -> bool:
    p = c.get("profile", {})
    career = c.get("career_history", [])
    skills = c.get("skills", [])

    total_months = sum(j.get("duration_months", 0) for j in career)
    stated_years = p.get("years_of_experience", 0)
    if stated_years > 0 and total_months > 0:
        if (total_months / 12) / stated_years > 2.5:
            return True

    expert_count = sum(1 for s in skills if s.get("proficiency") == "expert")
    total_endorsements = sum(s.get("endorsements", 0) for s in skills)
    if expert_count >= 5 and total_endorsements == 0:
        return True

    title = p.get("current_title", "").lower()
    if any(nt in title for nt in NON_TECH_HONEYPOT_TITLES):
        has_tech_evidence = any(
            kw in j.get("description", "").lower()
            for j in career
            for kw in TECH_EVIDENCE_KW
        )
        if not has_tech_evidence:
            return True

    return False


# ══════════════════════════════════════════════════════════════════════════════
# FAST TITLE PRE-FILTER
# ══════════════════════════════════════════════════════════════════════════════

def is_relevant(c: dict) -> bool:
    p = c.get("profile", {})
    combined = (p.get("current_title", "") + " " + p.get("headline", "")).lower()
    career_hint = ""
    if c.get("career_history"):
        career_hint = c["career_history"][0].get("description", "")[:200].lower()
    return (
        any(kw in combined for kw in TECH_TITLE_KW) or
        any(kw in career_hint for kw in TECH_EVIDENCE_KW)
    )


# ══════════════════════════════════════════════════════════════════════════════
# CANDIDATE TEXT FOR BI-ENCODER
# ══════════════════════════════════════════════════════════════════════════════

def build_candidate_text(c: dict) -> str:
    p = c.get("profile", {})
    sig = c.get("redrob_signals", {})

    parts = [
        p.get("current_title", ""),
        p.get("headline", ""),
        p.get("summary", "")[:200],
        "Skills: " + ", ".join(s["name"] for s in c.get("skills", [])[:12]),
    ]

    for job in c.get("career_history", [])[:3]:
        parts.append(
            f"{job.get('title', '')} at {job.get('company', '')}. "
            f"{job.get('description', '')[:150]}"
        )

    for proj in c.get("projects", [])[:2]:
        parts.append(proj.get("description", "")[:100])

    for ach in c.get("achievements", [])[:2]:
        parts.append(str(ach)[:80])

    if sig.get("open_to_work_flag"):
        parts.append("Open to work.")

    gh = sig.get("github_activity_score", -1)
    if gh > 30:
        parts.append(f"GitHub {gh}/100.")

    return " ".join(parts)[:600]


# ══════════════════════════════════════════════════════════════════════════════
# RULE SCORE
# ══════════════════════════════════════════════════════════════════════════════

def rule_score(c: dict, jd: dict) -> RuleScores:
    p = c.get("profile", {})
    sig = c.get("redrob_signals", {})
    career = c.get("career_history", [])
    skills = c.get("skills", [])

    eff_a = jd["_effective_tier_a"]
    yoe_min = jd["yoe_min"]
    yoe_max = jd["yoe_max"]

    tier_a_score = tier_b_score = penalty_score = total_weight = 0.0
    tier_a_hits = 0

    for sk in skills:
        name = sk["name"].lower()
        pw = PROFICIENCY_W.get(sk.get("proficiency", "beginner"), 0.25)
        end_t = min(1.0, (sk.get("endorsements", 0) + 1) / 10.0)
        dur_t = min(1.0, sk.get("duration_months", 0) / 24.0)
        w = pw * (0.4 + 0.3 * end_t + 0.3 * dur_t)
        total_weight += w

        in_a = name in eff_a
        in_b = (not in_a) and (name in SKILLS_TIER_B)
        in_p = name in SKILLS_PENALTY
        is_core = name in SKILLS_CORE_JD
        is_supporting = name in SKILLS_SUPPORTING

        if in_a:
            multiplier = 1.5 if is_core else 1.0
            tier_a_score += w * multiplier
            tier_a_hits += 1
        elif in_b:
            multiplier = 0.6 if is_supporting else 0.4
            tier_b_score += w * multiplier

        if in_p:
            penalty_score += w * 0.3

    skill_score = 0.0
    if total_weight > 0:
        raw = max(0.0, (tier_a_score + tier_b_score - penalty_score) / total_weight)
        skill_score = min(1.0, raw + min(0.20, tier_a_hits * 0.04))

    for skill_name, score_val in sig.get("skill_assessment_scores", {}).items():
        if skill_name.lower() in eff_a:
            skill_score = min(1.0, skill_score + (score_val / 100.0) * 0.05)

    title = p.get("current_title", "").lower()
    headline = p.get("headline", "").lower()
    combined = title + " " + headline

    all_career_text = " ".join(
        j.get("description", "") + " " +
        j.get("title", "") + " " +
        j.get("company", "")
        for j in career
    ).lower()

    all_project_text = " ".join(
        proj.get("description", "") + " " + proj.get("name", "")
        for proj in c.get("projects", [])
    ).lower()

    all_achievement_text = " ".join(
        str(a) for a in c.get("achievements", [])
    ).lower()

    all_responsibility_text = " ".join(
        " ".join(str(r) for r in j.get("responsibilities", []))
        for j in career
    ).lower()

    full_text = (
        all_career_text + " " +
        all_project_text + " " +
        all_achievement_text + " " +
        all_responsibility_text
    )

    content_fit_bonus = 0.0
    if RETRIEVAL_TITLE_RE.search(combined):
        retrieval_evidence = _count_prod_retrieval_hits(full_text)
        if retrieval_evidence >= 3:
            content_fit_bonus = 0.15
        elif retrieval_evidence >= 1:
            content_fit_bonus = 0.08

    if TITLE_FIT_RE.search(combined) and not ML_ADJACENT_TITLE_RE.search(title):
        title_score = 0.35
    elif ML_ADJACENT_TITLE_RE.search(title):
        title_score = 0.20
    else:
        title_score = 0.0

    title_score = min(0.35, title_score + content_fit_bonus)

    if TITLE_PENALTY_RE.search(title):
        title_score = max(0.0, title_score - 0.20)
    if WRONG_DOMAIN_RE.search(title):
        title_score = max(0.0, title_score - 0.30)
    if JUNIOR_TITLE_RE.search(title) and jd.get("seniority") in ("senior", "staff"):
        title_score = max(0.0, title_score - 0.15)

    if SENIOR_TITLE_RE.search(title):
        title_score = min(0.35, title_score + 0.05)

    yoe = p.get("years_of_experience", 0)
    if yoe_min <= yoe <= yoe_max:
        exp_score = 0.20
    elif (yoe_min - 1) <= yoe < yoe_min:
        exp_score = 0.08
    elif yoe > yoe_max:
        overage = yoe - yoe_max
        exp_score = max(-0.10, 0.08 - overage * 0.03)
    else:
        exp_score = -0.10

    svc_months = prod_months = total_months = 0
    prod_hits_product = 0
    prod_hits_any = 0
    prod_retrieval_hits = 0
    consulting_only = True
    jd_exact_signal = False
    product_co_months = 0

    for job in career:
        desc = job.get("description", "").lower()
        dur = job.get("duration_months", 0)
        total_months += dur

        co = job.get("company", "").lower()
        ind = job.get("industry", "").lower()

        is_svc = (
            (
                any(s in co for s in SERVICES_COMPANIES) or
                "it services" in ind or
                "consulting" in ind
            ) and not any(p2 in co for p2 in KNOWN_PRODUCT_COMPANIES)
        )

        is_product_co = _is_product_company_job(job) and not is_svc

        if is_svc:
            svc_months += dur
        else:
            prod_months += dur
            consulting_only = False

        if is_product_co:
            product_co_months += dur

        if _infer_production(desc):
            prod_hits_any += 1
            if is_product_co:
                prod_hits_product += 1

        if _count_prod_retrieval_hits(desc) >= 1:
            prod_retrieval_hits += 1

        if not jd_exact_signal and any(kw in desc for kw in JD_EXACT_KW):
            jd_exact_signal = True

    for proj in c.get("projects", []):
        proj_text = (proj.get("description", "") + " " + proj.get("name", "")).lower()
        if _infer_production(proj_text):
            prod_hits_any += 1
        if _count_prod_retrieval_hits(proj_text) >= 1:
            prod_retrieval_hits += 1

    if _infer_production(all_achievement_text):
        prod_hits_any += 1

    prod_ratio = (prod_months / total_months * 0.20) if total_months > 0 else 0.05

    product_co_ratio_bonus = (
        (product_co_months / total_months * 0.08) if total_months > 0 else 0.0
    )

    consult_pen = (
        -0.10 if (consulting_only and len(career) > 1 and jd.get("prefers_product"))
        else 0.0
    )

    if prod_hits_product >= 2:
        prod_bonus = 0.25
    elif prod_hits_product == 1:
        prod_bonus = 0.15
    elif prod_hits_any >= 2:
        prod_bonus = 0.10
    elif prod_hits_any == 1:
        prod_bonus = 0.05
    else:
        prod_bonus = 0.0

    if prod_retrieval_hits >= 3:
        prod_retrieval_bonus = 0.12
    elif prod_retrieval_hits >= 2:
        prod_retrieval_bonus = 0.08
    elif prod_retrieval_hits >= 1:
        prod_retrieval_bonus = 0.04
    else:
        prod_retrieval_bonus = 0.0

    jd_signal_bonus = 0.10 if jd_exact_signal else 0.0

    gh = sig.get("github_activity_score", -1)
    gh_bonus = (
        0.15 if gh > 50 else
        0.08 if gh > 20 else
        0.03 if gh > 0 else 0.0
    )

    acad_pen = _academic_penalty(full_text)
    progression_bonus = _career_progression_score(career)
    prod_eng_bonus = _product_engineering_bonus(full_text)

    career_score = min(1.0, max(0.0,
        title_score + exp_score + prod_ratio +
        consult_pen + prod_bonus + prod_retrieval_bonus +
        product_co_ratio_bonus + jd_signal_bonus + gh_bonus +
        acad_pen + progression_bonus + prod_eng_bonus
    ))

    if WRONG_DOMAIN_RE.search(title):
        career_score = min(career_score, 0.30)

    # ── Positive AI/ML corroboration gate ──────────────────────────────────
    ai_ml_title_fit = bool(TITLE_FIT_RE.search(combined) or RETRIEVAL_TITLE_RE.search(combined))
    has_verified_assessment = any(
        k.lower() in eff_a for k in sig.get("skill_assessment_scores", {})
    )
    has_corroborating_evidence = (
        ai_ml_title_fit or
        prod_hits_any >= 1 or
        prod_retrieval_hits >= 1 or
        jd_exact_signal or
        has_verified_assessment
    )
    if not has_corroborating_evidence:
        career_score = min(career_score, 0.20)

    c["_v3_signals"] = {
        "prod_retrieval_hits": prod_retrieval_hits,
        "prod_hits_any": prod_hits_any,
        "prod_hits_product": prod_hits_product,
        "jd_exact_signal": jd_exact_signal,
        "gh": gh,
        "product_co_months": product_co_months,
        "total_months": total_months,
        "prod_eng_bonus": prod_eng_bonus,
        "progression_bonus": progression_bonus,
        "full_text": full_text,
        "has_corroborating_evidence": has_corroborating_evidence,
    }

    country = p.get("country", "").lower()
    location = p.get("location", "").lower()
    relocate = sig.get("willing_to_relocate", False)
    target_country = (jd.get("target_country") or "").lower()

    in_target_city = any(city in location for city in jd["target_cities"])
    in_target_country = (not target_country) or (country == target_country)

    if in_target_city:
        loc_score = 1.0
    elif in_target_country and relocate:
        loc_score = 0.80
    elif in_target_country:
        loc_score = 0.60
    elif relocate:
        loc_score = 0.35
    else:
        loc_score = 0.15

    cand_mode = sig.get("preferred_work_mode", "flexible")
    if cand_mode == "flexible" or cand_mode == jd.get("work_mode", "hybrid"):
        loc_score = min(1.0, loc_score + 0.05)

    behav = 0.0

    days_inactive = days_since(parse_date(sig.get("last_active_date")))
    behav += (
        0.20 if days_inactive <= 14 else
        0.17 if days_inactive <= 30 else
        0.13 if days_inactive <= 60 else
        0.08 if days_inactive <= 90 else
        0.04 if days_inactive <= 180 else 0.01
    )

    if sig.get("open_to_work_flag"):
        behav += 0.12

    rr = sig.get("recruiter_response_rate", 0)
    behav += 0.12 if rr >= 0.7 else 0.08 if rr >= 0.4 else 0.04 if rr >= 0.2 else 0.0

    art = sig.get("avg_response_time_hours", 999)
    behav += 0.06 if art <= 4 else 0.04 if art <= 24 else 0.02 if art <= 72 else 0.0

    notice = sig.get("notice_period_days", 90)
    behav += (
        0.12 if notice <= 15 else
        0.09 if notice <= 30 else
        0.06 if notice <= 60 else
        0.03 if notice <= 90 else 0.01
    )

    icr = sig.get("interview_completion_rate", 0)
    behav += 0.08 if icr >= 0.8 else 0.05 if icr >= 0.6 else 0.02 if icr >= 0.4 else 0.0

    oar = sig.get("offer_acceptance_rate", -1)
    behav += (
        0.08 if oar >= 0.7 else
        0.05 if oar >= 0.4 else
        0.02 if oar >= 0.1 else
        0.03 if oar == -1 else 0.0
    )

    behav += sig.get("profile_completeness_score", 0) / 100.0 * 0.05

    saved = sig.get("saved_by_recruiters_30d", 0)
    behav += 0.05 if saved >= 10 else 0.03 if saved >= 5 else 0.01 if saved >= 2 else 0.0

    apps = sig.get("applications_submitted_30d", 0)
    behav += 0.03 if apps >= 3 else 0.01 if apps >= 1 else 0.0

    if sig.get("verified_email"):
        behav += 0.02
    if sig.get("verified_phone"):
        behav += 0.02
    if sig.get("linkedin_connected"):
        behav += 0.01

    conn = sig.get("connection_count", 0)
    behav += 0.02 if conn >= 300 else 0.01 if conn >= 100 else 0.0

    endorsements_received = sig.get("endorsements_received", 0)
    behav += (
        0.03 if endorsements_received >= 20 else
        0.02 if endorsements_received >= 10 else
        0.01 if endorsements_received >= 3 else 0.0
    )

    views = sig.get("profile_views_received_30d", 0)
    behav += 0.03 if views >= 10 else 0.02 if views >= 5 else 0.01 if views >= 2 else 0.0

    appearances = sig.get("search_appearance_30d", 0)
    behav += 0.02 if appearances >= 10 else 0.01 if appearances >= 3 else 0.0

    behav_score = min(1.0, max(0.0, behav))

    total = (
        W_SKILL * skill_score +
        W_CAREER * career_score +
        W_LOC * loc_score +
        W_BEHAV * behav_score
    )
    return RuleScores(
        total=total, skill=skill_score, career=career_score,
        loc=loc_score, behav=behav_score,
    )


# ══════════════════════════════════════════════════════════════════════════════
# TIE-BREAKING
# ══════════════════════════════════════════════════════════════════════════════

def tiebreak_key(item: tuple) -> tuple:
    hybrid, sem_n, rs, c = item
    sig = c.get("redrob_signals", {})
    sv = c.get("_v3_signals", {})
    total_months = max(sv.get("total_months", 1), 1)
    pc_ratio = sv.get("product_co_months", 0) / total_months
    notice = sig.get("notice_period_days", 90)
    rr = sig.get("recruiter_response_rate", 0)
    gh = sv.get("gh", -1)
    return (
        hybrid,
        sv.get("prod_retrieval_hits", 0),
        int(sv.get("jd_exact_signal", False)),
        round(pc_ratio, 2),
        sv.get("prod_hits_any", 0),
        round(sv.get("prod_eng_bonus", 0.0), 2),
        gh if gh > 0 else 0,
        rr,
        -notice,
    )


# ══════════════════════════════════════════════════════════════════════════════
# VALIDATION PASS
# ══════════════════════════════════════════════════════════════════════════════

def _validate_top20(top100: list[tuple]) -> list[str]:
    warnings: list[str] = []
    prev_score = 1.0
    for i, (hybrid, sem_n, rs, c) in enumerate(top100[:20]):
        if hybrid > prev_score + 1e-9:
            warnings.append(f"Rank {i+1}: score {hybrid:.4f} > prev {prev_score:.4f} — not sorted")
        prev_score = hybrid

        sv = c.get("_v3_signals", {})
        pr = sv.get("prod_retrieval_hits", 0)
        pa = sv.get("prod_hits_any", 0)
        cid = c.get("candidate_id", "?")

        if i < 10 and pr == 0 and pa == 0 and rs.skill < 0.40:
            warnings.append(
                f"Rank {i+1} (id={cid}): top-10 with 0 production evidence "
                f"and skill={rs.skill:.2f} — check scoring"
            )

    return warnings


# ══════════════════════════════════════════════════════════════════════════════
# REASONING BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def _find_production_company(c: dict) -> tuple[str, bool]:
    career = c.get("career_history", [])
    fallback_svc = ("", False)

    for job in career[:4]:
        desc = job.get("description", "").lower()
        co = job.get("company", "").lower()
        is_svc = (
            any(s in co for s in SERVICES_COMPANIES) and
            not any(p in co for p in KNOWN_PRODUCT_COMPANIES)
        )
        if _infer_production(desc):
            if not is_svc and _is_product_company_job(job):
                return job.get("company", ""), False
            elif not fallback_svc[0]:
                fallback_svc = (job.get("company", ""), True)

    for proj in c.get("projects", []):
        proj_text = (proj.get("description", "") + " " + proj.get("name", "")).lower()
        if _infer_production(proj_text) and not fallback_svc[0]:
            fallback_svc = (proj.get("name", "a project"), False)

    return fallback_svc


def _find_ops_signals(c: dict) -> list[str]:
    career = c.get("career_history", [])
    all_proj_text = " ".join(
        proj.get("description", "") + " " + proj.get("name", "")
        for proj in c.get("projects", [])
    ).lower()
    all_ach_text = " ".join(str(a) for a in c.get("achievements", [])).lower()

    found: list[str] = []
    for keyword, label in OPS_SIGNALS:
        hit = any(keyword in job.get("description", "").lower() for job in career)
        if not hit:
            hit = keyword in all_proj_text
        if not hit:
            hit = keyword in all_ach_text
        if hit and label not in found:
            found.append(label)
    return found


def _build_location_string(
    location: str, country: str, target_country: str,
    in_target: bool, is_intl: bool, relocate: bool,
) -> str:
    if in_target:
        return f"{location} (target city)"
    if country.lower() == target_country and relocate:
        return f"{location}, open to relocate"
    if country.lower() == target_country:
        return location
    suffix = ", open to relocate" if relocate else ""
    return f"{location}, {country}{suffix}"


def build_reasoning(c: dict, jd: dict, sem_n: float, rs: RuleScores, tier_label: str) -> str:
    p = c.get("profile", {})
    sig = c.get("redrob_signals", {})
    skills = c.get("skills", [])
    career = c.get("career_history", [])
    sv = c.get("_v3_signals", {})

    eff_a = jd["_effective_tier_a"]
    yoe_min = jd["yoe_min"]
    yoe_max = jd["yoe_max"]
    title = p.get("current_title", "")
    yoe = p.get("years_of_experience", 0)
    cid = c.get("candidate_id", "")

    matched_skills = sorted(
        [sk for sk in skills if sk["name"].lower() in eff_a],
        key=lambda x: x.get("endorsements", 0),
        reverse=True,
    )
    top_skills = [s["name"] for s in matched_skills[:3]]

    prod_company, prod_is_services = _find_production_company(c)
    ops_found = _find_ops_signals(c)

    best_assessment = ""
    best_val = -1.0
    for k, v in sig.get("skill_assessment_scores", {}).items():
        if k.lower() in eff_a and v > best_val:
            best_val = v
            best_assessment = f"{k} {v:.0f}/100"

    gh = sig.get("github_activity_score", -1)

    location = p.get("location", "")
    country = p.get("country", "")
    target_country = (jd.get("target_country") or "").lower()
    relocate = sig.get("willing_to_relocate", False)
    in_target = any(city in location.lower() for city in jd["target_cities"])
    is_intl = bool(target_country and country.lower() != target_country)

    loc_str = _build_location_string(
        location, country, target_country, in_target, is_intl, relocate
    )

    notice = sig.get("notice_period_days", 90)
    open_flag = "open to work" if sig.get("open_to_work_flag") else "not flagged open"
    inactive = days_since(parse_date(sig.get("last_active_date")))
    active_str = f"active {inactive}d ago" if inactive < 365 else "inactive >1yr"
    rr = sig.get("recruiter_response_rate", 0)

    if yoe_min <= yoe <= yoe_max:
        exp_note = f"{yoe:.0f}yr"
    elif yoe > yoe_max:
        exp_note = f"{yoe:.0f}yr (above {yoe_max}yr ceiling)"
    elif yoe == yoe_min - 1:
        exp_note = f"{yoe:.0f}yr (1yr below floor)"
    else:
        exp_note = f"{yoe:.0f}yr (outside {yoe_min}-{yoe_max}yr range)"

    title_lower = title.lower()
    is_junior_title = bool(JUNIOR_TITLE_RE.search(title_lower))
    jd_wants_senior = jd.get("seniority") in ("senior", "staff")
    is_non_ml_primary = any(t in title_lower for t in NON_ML_PRIMARY_TITLES)
    prod_retrieval_hits = sv.get("prod_retrieval_hits", 0)
    full_text = sv.get("full_text", "")

    action_verb = _get_action_verb(cid)
    notable_tech = _extract_notable_tech(full_text)
    tech_str = " + ".join(notable_tech[:2]) if notable_tech else ""

    score_tag = (
        f"[skills {rs.skill:.2f} · career {rs.career:.2f} · "
        f"loc {rs.loc:.2f} · avail {rs.behav:.2f}]"
    )
    max_body = 300 - len(score_tag) - 1

    no_prod_concern = not prod_company
    no_prod_offset_text = ""
    if no_prod_concern and tier_label == "strong":
        if rs.skill >= 0.55 or sem_n >= 0.65:
            no_prod_concern = False
            if prod_retrieval_hits >= 2:
                no_prod_offset_text = (
                    "Strong retrieval/search expertise offsets limited explicit production signals."
                )
            elif tech_str:
                no_prod_offset_text = (
                    f"Deep {tech_str} expertise offsets limited explicit production signals."
                )
            elif rs.skill >= 0.60:
                no_prod_offset_text = (
                    "Excellent JD skill match offsets limited explicit production signals."
                )
            else:
                no_prod_offset_text = (
                    "High semantic alignment to JD offsets limited explicit production signals."
                )

    if tier_label == "strong":
        parts = [f"{title}, {exp_note}"]
        skill_mention = tech_str if tech_str else (" + ".join(top_skills[:2]) if top_skills else "")

        if prod_company and skill_mention and not prod_is_services:
            parts.append(f"{action_verb} {skill_mention} in production at {prod_company}")
        elif prod_company and skill_mention and prod_is_services:
            parts.append(f"{action_verb} {skill_mention} at {prod_company} (IT services)")
        elif prod_company and not prod_is_services:
            parts.append(f"Production deployment at {prod_company}")
        elif skill_mention:
            parts.append(f"Core expertise: {skill_mention}")

        if ops_found:
            parts.append(f"Evidence: {', '.join(ops_found[:2])}")
        if best_assessment:
            parts.append(f"Assessed {best_assessment}")
        if gh > 50:
            parts.append(f"GitHub {gh}/100")
        parts.append(f"{loc_str}; {open_flag}, {active_str}")

        concerns: list[str] = []
        if not sv.get("has_corroborating_evidence", True):
            concerns.append(
                "Skills list matches JD keywords, but title/career history/assessments "
                "show no corroborating AI/ML evidence — verify before advancing."
            )
        if no_prod_offset_text:
            concerns.append(no_prod_offset_text)
        elif no_prod_concern:
            concerns.append("No confirmed production deployment.")
        if is_junior_title and jd_wants_senior:
            concerns.append(f"Title suggests junior level — verify seniority for {jd['seniority']} role.")
        elif is_non_ml_primary and jd_wants_senior:
            concerns.append(f"'{title}' is adjacent — verify ML depth for {jd['seniority']} role.")
        if rs.skill < 0.45:
            concerns.append(f"Low JD skill match ({rs.skill:.2f}) — verify technical depth.")
        if is_intl:
            concerns.append(f"International ({country}) — relocation or visa required.")
        if rr < 0.2:
            concerns.append(f"Low recruiter response ({rr:.0%}).")
        elif notice > 90:
            concerns.append(f"Long notice ({notice}d).")
        elif inactive > 180:
            concerns.append(f"Inactive {inactive}d — verify availability.")

        concern_str = (" " + " ".join(concerns[:2])) if concerns else ""
        body = ". ".join(parts) + "." + concern_str

    elif tier_label == "adjacent":
        parts = [f"{title}, {exp_note}"]
        skill_mention = tech_str if tech_str else (", ".join(top_skills) if top_skills else "")
        if skill_mention:
            parts.append(f"Skills: {skill_mention}")
        if prod_company and not prod_is_services:
            parts.append(f"Production at {prod_company}")
        elif prod_company and prod_is_services:
            parts.append(f"Deployed at {prod_company} (IT services)")
        if ops_found:
            parts.append(ops_found[0])
        if best_assessment:
            parts.append(f"Assessed {best_assessment}")
        parts.append(f"{loc_str}; {open_flag}, {active_str}")

        concerns: list[str] = []
        if not sv.get("has_corroborating_evidence", True):
            concerns.append("Skills list unverified by title/production/assessment evidence.")
        if no_prod_concern:
            concerns.append("No confirmed production deployment.")
        if is_junior_title and jd_wants_senior:
            concerns.append(f"Junior title for {jd['seniority']} role.")
        elif is_non_ml_primary:
            concerns.append("Adjacent title — ML depth unconfirmed.")
        if is_intl:
            concerns.append("International — relocation needed.")
        if rr < 0.2:
            concerns.append(f"Low recruiter response ({rr:.0%}).")
        elif notice > 90:
            concerns.append(f"Long notice ({notice}d).")
        elif inactive > 120:
            concerns.append(f"Inactive {inactive}d.")

        concern_str = (" " + " ".join(concerns[:2])) if concerns else ""
        body = ". ".join(parts) + "." + concern_str

    else:
        gaps: list[str] = []

        if not sv.get("has_corroborating_evidence", True):
            gaps.append("skills list unverified by title/production/assessment evidence")

        if no_prod_concern:
            gaps.append("no confirmed production deployment")
        elif prod_is_services:
            gaps.append(f"production only at IT services ({prod_company})")

        if is_junior_title and jd_wants_senior:
            gaps.append(f"junior title for {jd['seniority']} role")
        elif is_non_ml_primary:
            gaps.append(f"adjacent title ({title})")

        if is_intl:
            gaps.append(f"international ({country})")
        elif not in_target and not relocate and target_country and country.lower() == target_country:
            gaps.append(f"non-target city ({location}), not open to relocate")

        if yoe > yoe_max:
            gaps.append(f"{yoe:.0f}yr above {yoe_max}yr ceiling")
        elif yoe < yoe_min - 1:
            gaps.append(f"{yoe:.0f}yr below {yoe_min}yr floor")

        if rr < 0.3:
            gaps.append(f"low recruiter response ({rr:.0%})")
        if notice > 90:
            gaps.append(f"long notice ({notice}d)")
        if inactive > 90:
            gaps.append(f"inactive {inactive}d")
        if not sig.get("open_to_work_flag"):
            gaps.append("not flagged open to work")

        if rs.skill < 0.35:
            gaps.append(f"low JD skill overlap ({rs.skill:.2f})")
        elif not top_skills:
            gaps.append("limited JD skill overlap")

        if not gaps:
            if yoe > yoe_max:
                gaps.append(f"{yoe:.0f}yr above {yoe_max}yr ceiling")
            elif not in_target and not relocate:
                gaps.append(f"non-target city ({location})")
            elif rs.skill < 0.55:
                gaps.append(f"below-average JD skill match ({rs.skill:.2f})")
            else:
                gaps.append("below top-tier hybrid score threshold")

        skill_str = tech_str if tech_str else (", ".join(top_skills[:2]) if top_skills else "general ML background")
        body = (
            f"{title}, {exp_note}; skills: {skill_str}; {loc_str}. "
            f"Completing top 100; gaps: {'; '.join(gaps[:3])}."
        )

    body = _safe_truncate(body, max_body)
    return f"{body} {score_tag}"


# ══════════════════════════════════════════════════════════════════════════════
# DATA LOADER  (accepts any uploaded file — sniffs gzip/jsonl by content,
#               not by filename extension)
# ══════════════════════════════════════════════════════════════════════════════

def load_candidates(path_str: str, progress=None) -> list[dict]:
    """
    Loads a candidate file regardless of its filename/extension.

    HF Spaces / gr.File can hand back files with odd or missing extensions
    (e.g. a temp upload path). Instead of trusting `.gz`, this sniffs the
    first two bytes for the gzip magic number (1f 8b) and picks the opener
    accordingly, then parses each line as JSON.

    Speed/UX improvements for large files:
      - Uses orjson (C-based) for JSON parsing when available — falls back
        to stdlib json automatically if orjson isn't installed.
      - Reads/parses lines as bytes (skips Python text-decoding overhead).
      - Streams progress updates based on bytes read (uncompressed files;
        gzip files don't expose a reliable byte-progress cheaply, so those
        just show a periodic "N read" message instead).
    """
    path = Path(path_str)

    with open(path, "rb") as sniff:
        magic = sniff.read(2)
    is_gz = magic == b"\x1f\x8b"
    opener = gzip.open if is_gz else open

    try:
        import orjson
        loads = orjson.loads
    except ImportError:
        loads = json.loads

    out: list[dict] = []
    file_size = path.stat().st_size if not is_gz else None
    read_bytes = 0
    last_pct = -1

    with opener(path, "rb", buffering=4 * 1024 * 1024) as f:
        for raw_line in f:
            read_bytes += len(raw_line)
            if raw_line[:1] == b"{":
                try:
                    out.append(loads(raw_line))
                except ValueError:
                    pass

            if progress is not None:
                if file_size:
                    pct = int(read_bytes / file_size * 100)
                    if pct != last_pct and pct % 5 == 0:
                        progress(0.08 + 0.07 * (pct / 100), desc=f"Loading candidates... {pct}%")
                        last_pct = pct
                elif len(out) % 5000 == 0 and len(out) > 0:
                    progress(0.10, desc=f"Loading candidates... {len(out):,} read")

    if not out:
        raise gr.Error(
            "Couldn't find any JSON candidate records in that file. "
            "Make sure it's a JSONL (one JSON object per line) or gzipped JSONL file."
        )

    return out


# ══════════════════════════════════════════════════════════════════════════════
# MODEL CACHE (loaded once per Space instance, reused across runs)
# ══════════════════════════════════════════════════════════════════════════════

_MODEL_CACHE: dict[str, SentenceTransformer] = {}


def get_model(model_name: str = BIENCODER_MODEL) -> SentenceTransformer:
    if model_name not in _MODEL_CACHE:
        _MODEL_CACHE[model_name] = SentenceTransformer(model_name)
    return _MODEL_CACHE[model_name]


def _format_jd_summary(jd: dict, jd_text: str) -> str:
    """Human-readable summary of what parse_jd() extracted from the uploaded
    JD. Entirely derived from the parsed jd dict — nothing here is specific
    to any particular job description; it just renders whatever was detected."""
    skills_list = sorted(jd["jd_skills"])
    cities = ", ".join(city.title() for city in jd["target_cities"]) or "Not specified"
    country = jd["target_country"].title() if jd["target_country"] else "Not specified"
    product_pref = (
        "Yes — JD language favors product-company background over services/consulting"
        if jd["prefers_product"] else "No strong signal detected"
    )
    skills_str = ", ".join(skills_list) if skills_list else "No reference skills matched in the JD text"

    return (
        "#### 🎯 What we understood from this job description\n\n"
        "| Attribute | Detected value |\n"
        "|---|---|\n"
        f"| **Experience range** | {jd['yoe_min']}–{jd['yoe_max']} years |\n"
        f"| **Seniority level** | {jd['seniority'].title()} |\n"
        f"| **Work mode** | {jd['work_mode'].title()} |\n"
        f"| **Target cities** | {cities} |\n"
        f"| **Target country** | {country} |\n"
        f"| **Product-company preference** | {product_pref} |\n"
        f"| **JD skills recognized** | {len(skills_list)} |\n"
        f"| **JD length** | {len(jd_text.split()):,} words |\n\n"
        f"**Skills recognized in this JD:** {skills_str}"
    )


def analyze_jd(jd_file, jd_text_box) -> str:
    """Fast, standalone JD parse — lets the user confirm what the ranker
    understood before committing to a full candidate-ranking run."""
    if not (jd_text_box or "").strip() and jd_file is None:
        raise gr.Error("Upload a JD file or paste JD text first.")
    log: list[str] = []
    if (jd_text_box or "").strip():
        jd_text = jd_text_box
    else:
        jd_text = load_jd_from_path(jd_file.name, log)
    jd = parse_jd(jd_text)
    return _format_jd_summary(jd, jd_text)


# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE (adapted from rank.py main())
# ══════════════════════════════════════════════════════════════════════════════

def run_pipeline(
    candidates_file,
    jd_file,
    jd_text_box,
    prefilter,
    preview_n,
    strong_pct,
    adjacent_pct,
    progress=gr.Progress(track_tqdm=True),
):
    if candidates_file is None:
        raise gr.Error("Please upload a candidates file (JSONL or gzipped JSONL).")
    if jd_file is None and not (jd_text_box or "").strip():
        raise gr.Error("Please upload a JD file or paste the JD text.")

    t0 = time.time()
    log: list[str] = []

    progress(0.02, desc="Parsing job description...")
    if (jd_text_box or "").strip():
        jd_text = jd_text_box
        log.append(f"Using pasted JD text ({len(jd_text.split())} words)")
    else:
        jd_text = load_jd_from_path(jd_file.name, log)

    jd = parse_jd(jd_text)
    log.append(f"YoE range     : {jd['yoe_min']}-{jd['yoe_max']} yrs")
    log.append(f"Target cities : {jd['target_cities'] or 'not specified'}")
    log.append(f"Target country: {jd['target_country'] or 'not specified'}")
    log.append(f"JD skills     : {len(jd['jd_skills'])} detected")
    log.append(f"Seniority     : {jd['seniority']}  |  Work mode: {jd['work_mode']}  |  Prefers product: {jd['prefers_product']}")
    log.append(f"Weights       : skill {W_SKILL} · career {W_CAREER} · loc {W_LOC} · avail {W_BEHAV}")

    progress(0.08, desc="Loading candidates...")
    all_candidates = load_candidates(candidates_file.name, progress=progress)
    log.append(f"Loaded {len(all_candidates):,} candidates")

    clean = [c for c in all_candidates if not is_honeypot(c)]
    log.append(f"{len(clean):,} clean ({len(all_candidates) - len(clean):,} honeypots removed)")

    progress(0.15, desc="Fast pre-filter + rule scoring...")
    relevant = [c for c in clean if is_relevant(c)]
    log.append(f"Fast filter: {len(clean):,} -> {len(relevant):,} relevant")

    def cheap_score(c: dict) -> float:
        p = c.get("profile", {})
        t = (p.get("current_title", "") + " " + p.get("headline", "")).lower()
        yoe = p.get("years_of_experience", 0)
        fit = (1.0 if TITLE_FIT_RE.search(t)
               else 0.5 if ML_ADJACENT_TITLE_RE.search(t)
               else 0.0)
        pen = 0.3 if TITLE_PENALTY_RE.search(t) else 0.0
        wr_pen = 0.4 if WRONG_DOMAIN_RE.search(t) else 0.0
        jr_pen = (0.2 if JUNIOR_TITLE_RE.search(t) and jd.get("seniority") in ("senior", "staff") else 0.0)
        in_rng = 1.0 if jd["yoe_min"] <= yoe <= jd["yoe_max"] else 0.0
        return fit - pen - wr_pen - jr_pen + in_rng * 0.5

    relevant.sort(key=cheap_score, reverse=True)
    stage_b_pool = relevant[:3000]
    log.append(f"Stage-A filter: {len(relevant):,} -> {len(stage_b_pool):,} for full scoring")

    progress(0.30, desc="Rule scoring candidates...")
    scored_pairs: list[tuple[RuleScores, dict]] = [
        (rule_score(c, jd), c) for c in stage_b_pool
    ]
    scored_pairs.sort(key=lambda x: -x[0].total)

    prefilter = int(prefilter)
    pool = [c for rs, c in scored_pairs[:prefilter]]
    pool_scores = {c["candidate_id"]: rs for rs, c in scored_pairs}

    if scored_pairs:
        best_rs = scored_pairs[0][0].total
        worst_rs = scored_pairs[min(prefilter - 1, len(scored_pairs) - 1)][0].total
        log.append(f"Top {prefilter} score range: {worst_rs:.3f} - {best_rs:.3f}")

    progress(0.45, desc=f"Loading bi-encoder ({BIENCODER_MODEL})...")
    bi_model = get_model(BIENCODER_MODEL)

    jd_emb = bi_model.encode(
        [jd_text], normalize_embeddings=True, convert_to_numpy=True, show_progress_bar=False,
    )

    progress(0.55, desc=f"Encoding {len(pool)} candidate profiles...")
    texts = [build_candidate_text(c) for c in pool]
    cand_embs = bi_model.encode(
        texts, batch_size=32, normalize_embeddings=True,
        convert_to_numpy=True, show_progress_bar=False,
    )

    dim = cand_embs.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(cand_embs.astype(np.float32))
    distances, indices = index.search(jd_emb.astype(np.float32), len(pool))

    sem_raw = distances[0]
    s_min, s_max = float(sem_raw.min()), float(sem_raw.max())
    sem_norm = (sem_raw - s_min) / (s_max - s_min + 1e-9)
    log.append(f"Semantic range: {s_min:.3f} - {s_max:.3f}")

    progress(0.80, desc="Hybrid scoring, validation and writing top 100...")
    results: list[tuple[float, float, RuleScores, dict]] = []
    for rank_i, idx in enumerate(indices[0]):
        c = pool[idx]
        sem_n = float(sem_norm[rank_i])
        rs = pool_scores[c["candidate_id"]]
        hybrid = W_SEMANTIC * sem_n + W_RULES * rs.total
        results.append((hybrid, sem_n, rs, c))

    results.sort(key=tiebreak_key, reverse=True)

    yoe_min = jd["yoe_min"]
    yoe_max = jd["yoe_max"]

    def yoe_in_range(c: dict) -> bool:
        return yoe_min <= c["profile"].get("years_of_experience", 0) <= yoe_max

    in_range = [r for r in results if yoe_in_range(r[3])]
    out_range = [r for r in results if not yoe_in_range(r[3])]

    final = in_range[:100]
    if len(final) < 100:
        pad = 100 - len(final)
        final += out_range[:pad]
        log.append(f"Note: padded {pad} out-of-range candidates")

    final.sort(key=tiebreak_key, reverse=True)
    top100 = final[:100]

    if not top100:
        raise gr.Error("No candidates matched. Check the candidates file / JD content.")

    top100_scores = np.array([r[0] for r in top100])
    strong_thresh = float(np.percentile(top100_scores, float(strong_pct)))
    adjacent_thresh = float(np.percentile(top100_scores, float(adjacent_pct)))
    log.append(
        f"Tier thresholds (p{strong_pct:.0f}/p{adjacent_pct:.0f}): "
        f"strong >{strong_thresh:.4f}, adjacent >{adjacent_thresh:.4f}, filler <={adjacent_thresh:.4f}"
    )

    warnings = _validate_top20(top100)
    if warnings:
        log.append(f"[VALIDATION] {len(warnings)} warning(s):")
        log.extend(f"  ! {w}" for w in warnings)
    else:
        log.append("[VALIDATION] Top-20 passed all checks.")

    rows: list[dict] = []
    n_strong = n_adjacent = n_filler = 0
    for rank_idx, (hybrid, sem_n, rs, c) in enumerate(top100):
        tier_label = (
            "strong" if hybrid >= strong_thresh else
            "adjacent" if hybrid >= adjacent_thresh else
            "filler"
        )
        if tier_label == "strong":
            n_strong += 1
        elif tier_label == "adjacent":
            n_adjacent += 1
        else:
            n_filler += 1
        rows.append({
            "candidate_id": c["candidate_id"],
            "rank": rank_idx + 1,
            "score": round(hybrid, 6),
            "tier": tier_label,
            "reasoning": build_reasoning(c, jd, sem_n, rs, tier_label),
        })

    out_dir = Path(tempfile.mkdtemp(prefix="redrob_"))
    out_path = out_dir / "submission.csv"
    with open(out_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["candidate_id", "rank", "score", "tier", "reasoning"])
        writer.writeheader()
        writer.writerows(rows)

    elapsed = time.time() - t0
    log.append(f"Submission written -> {out_path.name}")
    log.append(f"Total runtime: {elapsed:.1f}s ({elapsed / 60:.1f} min)")

    df = pd.DataFrame(rows)
    preview_n = max(1, min(int(preview_n), 100))
    preview_df = df.head(preview_n)

    stats_html = _format_stats_html(
        total=len(all_candidates),
        honeypots=len(all_candidates) - len(clean),
        relevant=len(relevant),
        n_strong=n_strong,
        n_adjacent=n_adjacent,
        n_filler=n_filler,
        elapsed=elapsed,
    )

    progress(1.0, desc="Done.")
    return preview_df, str(out_path), "\n".join(log), stats_html


def _format_stats_html(total, honeypots, relevant, n_strong, n_adjacent, n_filler, elapsed) -> str:
    stats = [
        ("📥", str(total), "Candidates loaded"),
        ("🛡️", str(honeypots), "Honeypots removed"),
        ("🎯", str(relevant), "Passed relevance filter"),
        ("🟢", str(n_strong), "Strong-tier matches"),
        ("🟡", str(n_adjacent), "Adjacent-tier matches"),
        ("⚪", str(n_filler), "Filler-tier matches"),
        ("⏱️", f"{elapsed:.1f}s", "Total runtime"),
    ]
    cards = "".join(
        f"""
        <div class="stat-card">
            <div class="stat-icon">{icon}</div>
            <div class="stat-value">{value}</div>
            <div class="stat-label">{label}</div>
        </div>
        """
        for icon, value, label in stats
    )
    return f'<div class="stats-grid">{cards}</div>'


# ══════════════════════════════════════════════════════════════════════════════
# GRADIO UI — REDMIND polish pass (custom theme + animations)
# ══════════════════════════════════════════════════════════════════════════════

CUSTOM_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

:root {
    --redrob-primary: #6D28D9;
    --redrob-primary-dark: #4C1D95;
    --redrob-accent: #F59E0B;
}

.gradio-container {
    max-width: 1400px !important;
    margin: 0 auto !important;
    font-family: 'Inter', 'Segoe UI', sans-serif !important;
    animation: fadeInPage 0.6s ease both;
}

@keyframes fadeInPage {
    from { opacity: 0; transform: translateY(8px); }
    to { opacity: 1; transform: translateY(0); }
}

@keyframes floatIn {
    from { opacity: 0; transform: translateY(14px); }
    to { opacity: 1; transform: translateY(0); }
}

@keyframes shimmer {
    0% { background-position: -400px 0; }
    100% { background-position: 400px 0; }
}

@keyframes pulseGlow {
    0%, 100% { box-shadow: 0 6px 18px rgba(109, 40, 217, 0.35); }
    50% { box-shadow: 0 10px 28px rgba(168, 85, 247, 0.55); }
}

@keyframes popIn {
    0% { opacity: 0; transform: scale(0.85); }
    70% { transform: scale(1.03); }
    100% { opacity: 1; transform: scale(1); }
}

@keyframes badgePulse {
    0%, 100% { transform: scale(1); }
    50% { transform: scale(1.06); }
}

/* Hero header */
#hero-header {
    position: relative;
    overflow: hidden;
    background: linear-gradient(120deg, #6D28D9 0%, #7C3AED 45%, #A855F7 100%);
    border-radius: 18px;
    padding: 34px 36px;
    margin-bottom: 22px;
    box-shadow: 0 10px 30px rgba(109, 40, 217, 0.25);
    animation: floatIn 0.7s ease both;
}
#hero-header::after {
    content: "";
    position: absolute;
    top: -60%;
    left: -20%;
    width: 60%;
    height: 220%;
    background: rgba(255,255,255,0.08);
    transform: rotate(20deg);
    animation: shimmer 5s linear infinite;
}
#hero-header h1 {
    color: #ffffff !important;
    font-size: 2.15rem !important;
    font-weight: 800 !important;
    margin-bottom: 4px !important;
}
#hero-header p {
    color: #EDE9FE !important;
    font-size: 1rem !important;
    margin: 0 !important;
    max-width: 760px;
}
#hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.18);
    color: #fff;
    border: 1px solid rgba(255,255,255,0.35);
    border-radius: 999px;
    padding: 4px 14px;
    font-size: 0.75rem;
    font-weight: 600;
    letter-spacing: 0.04em;
    margin-bottom: 10px;
    animation: badgePulse 2.4s ease-in-out infinite;
}

/* Section cards */
.redrob-card {
    background: #ffffff;
    border: 1px solid #ECECF3;
    border-radius: 16px;
    padding: 20px !important;
    box-shadow: 0 2px 10px rgba(20, 20, 43, 0.04);
    animation: floatIn 0.6s ease both;
    transition: box-shadow 0.25s ease, transform 0.25s ease;
}
.redrob-card:hover {
    box-shadow: 0 8px 24px rgba(109, 40, 217, 0.10);
    transform: translateY(-2px);
}

/* Section labels */
.section-label {
    font-weight: 700;
    font-size: 0.95rem;
    color: #4C1D95;
    margin-bottom: 6px;
    display: flex;
    align-items: center;
    gap: 6px;
}

/* Run button */
#run-btn {
    background: linear-gradient(120deg, #6D28D9, #A855F7) !important;
    border: none !important;
    color: #fff !important;
    font-weight: 700 !important;
    font-size: 1.05rem !important;
    border-radius: 12px !important;
    padding: 14px 0 !important;
    animation: pulseGlow 2.6s ease-in-out infinite;
    transition: transform 0.15s ease !important;
}
#run-btn:hover {
    transform: translateY(-2px) scale(1.01);
}
#run-btn:active {
    transform: translateY(0) scale(0.99);
}

/* Analyze button */
#analyze-btn {
    transition: transform 0.15s ease !important;
}
#analyze-btn:hover {
    transform: translateY(-1px);
}

/* Stats grid (post-run summary cards) */
.stats-grid {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(120px, 1fr));
    gap: 12px;
    margin-top: 4px;
}
.stat-card {
    background: linear-gradient(160deg, #FAFAFF 0%, #F3EEFE 100%);
    border: 1px solid #E9E2FB;
    border-radius: 14px;
    padding: 14px 10px;
    text-align: center;
    animation: popIn 0.5s ease both;
}
.stat-icon {
    font-size: 1.3rem;
    margin-bottom: 2px;
}
.stat-value {
    font-size: 1.4rem;
    font-weight: 800;
    color: var(--redrob-primary-dark);
    line-height: 1.1;
}
.stat-label {
    font-size: 0.72rem;
    color: #6B7280;
    font-weight: 600;
    margin-top: 2px;
}

/* Results table */
#results-table table {
    border-radius: 12px !important;
    overflow: hidden !important;
}
#results-table {
    animation: floatIn 0.5s ease both;
}

/* Footer */
#redmind-footer {
    text-align: center;
    color: #9CA3AF;
    font-size: 0.8rem;
    padding: 20px 0 8px 0;
    animation: floatIn 0.8s ease both;
}
"""

THEME = gr.themes.Soft(
    primary_hue=gr.themes.colors.purple,
    secondary_hue=gr.themes.colors.amber,
    neutral_hue=gr.themes.colors.slate,
    font=[gr.themes.GoogleFont("Inter"), "sans-serif"],
).set(
    button_primary_background_fill="linear-gradient(120deg, #6D28D9, #A855F7)",
    button_primary_background_fill_hover="linear-gradient(120deg, #5B21B6, #9333EA)",
    block_radius="16px",
    block_shadow="0 2px 10px rgba(20,20,43,0.05)",
)

with gr.Blocks(
    title="Redrob Ranker — Intelligent Candidate Discovery & Ranking",
    theme=THEME,
    css=CUSTOM_CSS,
) as demo:

    # ── Hero header ─────────────────────────────────────────────────────
    gr.HTML(
        """
        <div id="hero-header">
            <div id="hero-badge">REDMIND · India.RUNS 2026</div>
            <h1>🎯 Redrob Ranker</h1>
            <p>Hybrid semantic + rule-based candidate ranking engine — upload a candidate
            pool and a job description to generate a ranked top-100 shortlist with
            transparent, per-candidate reasoning.</p>
        </div>
        """
    )

    with gr.Row(equal_height=False):
        # ── Left: inputs ────────────────────────────────────────────────
        with gr.Column(scale=1):
            with gr.Group(elem_classes="redrob-card"):
                gr.Markdown("### 📥 Step 1 — Upload your data")

                gr.HTML('<div class="section-label">👥 Candidate pool</div>')
                candidates_file = gr.File(
                    label="Any file — JSONL / JSONL.GZ recommended, format is auto-detected",
                )

                gr.HTML('<div class="section-label">📄 Job description</div>')
                jd_file = gr.File(
                    label="Any file (optional if pasting below) — .md/.txt/.docx work best",
                )
                jd_text_box = gr.Textbox(
                    label="Or paste job description text",
                    lines=7,
                    placeholder="Paste the job description here...",
                )

                analyze_btn = gr.Button("🔍 Preview JD Understanding", size="sm", elem_id="analyze-btn")
                jd_summary_box = gr.Markdown(visible=True)

            with gr.Group(elem_classes="redrob-card"):
                with gr.Accordion("⚙️ Advanced settings", open=False):
                    prefilter = gr.Slider(
                        label="Rule-scoring prefilter pool size",
                        minimum=50, maximum=1000, step=50, value=300,
                    )
                    preview_n = gr.Slider(
                        label="Rows to preview in the table",
                        minimum=10, maximum=100, step=10, value=100,
                    )
                    strong_pct = gr.Slider(
                        label="Strong-tier percentile threshold",
                        minimum=0, maximum=100, step=5, value=60,
                    )
                    adjacent_pct = gr.Slider(
                        label="Adjacent-tier percentile threshold",
                        minimum=0, maximum=100, step=5, value=30,
                    )

            run_btn = gr.Button("🚀 Run Ranking", elem_id="run-btn", size="lg")

        # ── Right: outputs ──────────────────────────────────────────────
        with gr.Column(scale=2):
            with gr.Group(elem_classes="redrob-card"):
                gr.Markdown("### 📊 Run Summary")
                stats_output = gr.HTML(
                    '<div class="stats-grid"></div>'
                )

            with gr.Group(elem_classes="redrob-card"):
                gr.Markdown("### 🏆 Top Candidates")
                results_table = gr.Dataframe(
                    label="",
                    headers=["candidate_id", "rank", "score", "tier", "reasoning"],
                    wrap=True,
                    elem_id="results-table",
                )
                with gr.Row():
                    csv_output = gr.File(label="⬇️ Download submission.csv")

            with gr.Group(elem_classes="redrob-card"):
                gr.Markdown("### 📋 Run log")
                log_output = gr.Textbox(label="", lines=14, max_lines=30, show_label=False)

    gr.HTML('<div id="redmind-footer">Built by Team REDMIND · Redrob India.RUNS Data & AI 2026</div>')

    analyze_btn.click(
        fn=analyze_jd,
        inputs=[jd_file, jd_text_box],
        outputs=[jd_summary_box],
    )

    run_btn.click(
        fn=run_pipeline,
        inputs=[candidates_file, jd_file, jd_text_box, prefilter, preview_n, strong_pct, adjacent_pct],
        outputs=[results_table, csv_output, log_output, stats_output],
    )

if __name__ == "__main__":
    demo.queue().launch(show_api=False)